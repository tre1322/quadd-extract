"""
Simple Transformer - LLM-based document transformation with vision + OCR.

New architecture:
1. Convert PDF pages to images (for structure/layout)
2. Extract accurate OCR text at 300 DPI (for correct names/values)
3. Store example images + OCR text + desired output
4. For extraction: Send Claude BOTH images AND OCR text
5. Claude sees structure from images, gets accurate text from OCR

Generic approach - works for ANY PDF with image-rendered text:
- Honor rolls, legal notices, sports, etc.
- No document-type-specific code
- More robust PDF text extraction (charter compliance)

No complex column mapping, anchors, or regions. Just examples + vision + OCR.
"""
from __future__ import annotations

import logging
from typing import Optional, List, Dict
import pymupdf  # PyMuPDF for PDF to image conversion
import anthropic
import os
import base64
from io import BytesIO
from PIL import Image
import pytesseract

logger = logging.getLogger(__name__)


class SimpleTransformer:
    """
    Simple document transformer using Claude vision + OCR + examples.

    No complex IR, no column mapping, no anchors.
    Just: PDF images + OCR text + example → Claude → formatted output.

    Generic approach - works for ANY document type.
    """

    def __init__(self, api_key: Optional[str] = None):
        """Initialize transformer with Anthropic API key."""
        self.api_key = api_key or os.getenv("ANTHROPIC_API_KEY")
        if not self.api_key:
            raise ValueError("ANTHROPIC_API_KEY not set")

        self.client = anthropic.Anthropic(api_key=self.api_key)
        self.examples = {}  # processor_id -> (input_images, ocr_text, output_text)

    def detect_file_type(self, file_bytes: bytes, filename: str = None) -> str:
        """
        Detect file type from bytes and/or filename.

        Args:
            file_bytes: File content as bytes
            filename: Optional filename with extension

        Returns:
            File type: 'pdf', 'docx', 'doc', 'xlsx', 'xls', 'csv', 'txt', 'image'
        """
        # Check filename extension first
        if filename:
            ext = filename.lower().split('.')[-1]
            if ext == 'pdf':
                return 'pdf'
            elif ext in ['docx']:
                return 'docx'
            elif ext in ['doc']:
                return 'doc'
            elif ext in ['xlsx']:
                return 'xlsx'
            elif ext in ['xls']:
                return 'xls'
            elif ext == 'csv':
                return 'csv'
            elif ext == 'txt':
                return 'txt'
            elif ext in ['png', 'jpg', 'jpeg', 'gif', 'bmp']:
                return 'image'

        # Check magic bytes as fallback
        if file_bytes[:4] == b'%PDF':
            return 'pdf'
        elif file_bytes[:2] == b'PK':  # ZIP-based formats (docx, xlsx)
            if b'word/' in file_bytes[:1000]:
                return 'docx'
            elif b'xl/' in file_bytes[:1000]:
                return 'xlsx'
        elif file_bytes[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1':  # OLE format (doc, xls)
            return 'doc'  # Could be xls too, but we'll default to doc
        elif file_bytes[:2] in [b'\xff\xd8', b'\x89P']:  # JPEG or PNG
            return 'image'

        # Default to text if unknown
        return 'txt'

    def extract_text_from_txt(self, file_bytes: bytes) -> str:
        """Extract text from .txt file."""
        try:
            return file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            # Try other encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    return file_bytes.decode(encoding)
                except:
                    continue
            return file_bytes.decode('utf-8', errors='ignore')

    def extract_text_from_docx(self, file_bytes: bytes) -> str:
        """Extract text from .docx file."""
        from docx import Document
        from io import BytesIO

        doc = Document(BytesIO(file_bytes))
        text_parts = []

        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            text_parts.append(paragraph.text)

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = '\t'.join(cell.text for cell in row.cells)
                text_parts.append(row_text)

        return '\n'.join(text_parts)

    def extract_text_from_xlsx(self, file_bytes: bytes) -> str:
        """Extract text from .xlsx file."""
        import openpyxl
        from io import BytesIO

        wb = openpyxl.load_workbook(BytesIO(file_bytes), data_only=True)
        text_parts = []

        for sheet in wb.worksheets:
            text_parts.append(f"Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                row_text = '\t'.join(str(cell) if cell is not None else '' for cell in row)
                if row_text.strip():
                    text_parts.append(row_text)
            text_parts.append('')  # Blank line between sheets

        return '\n'.join(text_parts)

    def extract_text_from_csv(self, file_bytes: bytes) -> str:
        """Extract text from .csv file."""
        import pandas as pd
        from io import StringIO

        # Try to decode
        text = self.extract_text_from_txt(file_bytes)

        # Parse CSV
        try:
            df = pd.read_csv(StringIO(text))
            return df.to_string(index=False)
        except:
            # If pandas fails, return raw text
            return text

    def extract_text_from_xls(self, file_bytes: bytes) -> str:
        """Extract text from .xls file (old Excel format)."""
        import pandas as pd
        from io import BytesIO

        try:
            # Read all sheets
            excel_file = pd.ExcelFile(BytesIO(file_bytes))
            text_parts = []

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                text_parts.append(f"Sheet: {sheet_name}")
                text_parts.append(df.to_string(index=False))
                text_parts.append('')

            return '\n'.join(text_parts)
        except:
            # Fallback: try to extract as binary
            return self.extract_text_from_txt(file_bytes)

    def image_to_base64(self, image_bytes: bytes) -> str:
        """Convert image bytes to base64."""
        return base64.b64encode(image_bytes).decode('utf-8')

    def extract_from_image(self, image_bytes: bytes) -> tuple[List[str], str]:
        """
        Extract from image file (same as PDF: images for vision + OCR for text).

        Args:
            image_bytes: Image file bytes

        Returns:
            Tuple of (base64_images, ocr_text)
        """
        # Convert to base64 for vision
        b64_image = self.image_to_base64(image_bytes)

        # Extract OCR text
        image = Image.open(BytesIO(image_bytes))
        ocr_text = pytesseract.image_to_string(image)

        logger.info(f"Extracted image: {len(ocr_text)} chars OCR")

        return [b64_image], ocr_text

    def extract_content_from_file(self, file_bytes: bytes, filename: str = None) -> tuple[List[str], str, str]:
        """
        Extract content from any supported file type.

        Args:
            file_bytes: File content as bytes
            filename: Optional filename for type detection

        Returns:
            Tuple of (images_b64, text_content, file_type)
            - images_b64: List of base64 images (for PDF/images, empty for text files)
            - text_content: Extracted text
            - file_type: Detected file type
        """
        file_type = self.detect_file_type(file_bytes, filename)
        logger.info(f"Detected file type: {file_type}")

        if file_type == 'pdf':
            # PDF: images + OCR
            images = self.pdf_to_images(file_bytes, dpi=150)
            ocr_text = self.extract_ocr_text(file_bytes, dpi=300)
            return images, ocr_text, file_type

        elif file_type == 'image':
            # Image: convert to b64 + OCR
            images, ocr_text = self.extract_from_image(file_bytes)
            return images, ocr_text, file_type

        elif file_type == 'docx':
            # Word: text only, no images
            text = self.extract_text_from_docx(file_bytes)
            return [], text, file_type

        elif file_type == 'xlsx':
            # Excel: text only
            text = self.extract_text_from_xlsx(file_bytes)
            return [], text, file_type

        elif file_type == 'xls':
            # Old Excel: text only
            text = self.extract_text_from_xls(file_bytes)
            return [], text, file_type

        elif file_type == 'csv':
            # CSV: text only
            text = self.extract_text_from_csv(file_bytes)
            return [], text, file_type

        elif file_type == 'txt':
            # Plain text
            text = self.extract_text_from_txt(file_bytes)
            return [], text, file_type

        else:
            # Unknown: try text extraction
            logger.warning(f"Unknown file type, attempting text extraction")
            text = self.extract_text_from_txt(file_bytes)
            return [], text, 'txt'

    def pdf_to_images(self, pdf_bytes: bytes, dpi: int = 150) -> List[str]:
        """
        Convert PDF pages to base64-encoded PNG images.

        Args:
            pdf_bytes: PDF file bytes
            dpi: Resolution for rendering (150 for display, 300 for OCR)

        Returns:
            List of base64-encoded PNG images (one per page)
        """
        doc = pymupdf.open(stream=pdf_bytes, filetype="pdf")
        images_b64 = []

        for page_num in range(len(doc)):
            page = doc[page_num]

            # Render page to image (PNG)
            # zoom = dpi / 72 (72 is default DPI)
            zoom = dpi / 72
            mat = pymupdf.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)

            # Convert to PNG bytes
            png_bytes = pix.tobytes("png")

            # Encode to base64
            b64_image = base64.b64encode(png_bytes).decode('utf-8')
            images_b64.append(b64_image)

            logger.info(f"Converted page {page_num + 1} to image ({len(png_bytes)} bytes, {dpi} DPI)")

        doc.close()

        logger.info(f"Converted {len(images_b64)} pages to images at {dpi} DPI")
        return images_b64

    def extract_ocr_text(self, pdf_bytes: bytes, dpi: int = 300) -> str:
        """
        Extract text from PDF using Tesseract OCR at high resolution.

        This provides accurate text extraction for image-rendered text
        (player names, numbers, etc.) that PyMuPDF's get_text() misses.

        Generic approach - works for ANY document type.

        Args:
            pdf_bytes: PDF file bytes
            dpi: Resolution for OCR (300 DPI recommended for accuracy)

        Returns:
            OCR-extracted text from all pages
        """
        doc = pymupdf.open(stream=pdf_bytes, filetype="pdf")
        ocr_text_parts = []

        for page_num in range(len(doc)):
            page = doc[page_num]

            # Render page at high DPI for accurate OCR
            zoom = dpi / 72
            mat = pymupdf.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)

            # Convert to PIL Image
            img_bytes = pix.tobytes("png")
            img = Image.open(BytesIO(img_bytes))

            # Run Tesseract OCR
            ocr_text = pytesseract.image_to_string(img)

            ocr_text_parts.append(f"--- Page {page_num + 1} (OCR Text) ---\n{ocr_text}")

            logger.info(f"OCR extracted {len(ocr_text)} chars from page {page_num + 1}")

        doc.close()

        full_ocr_text = "\n\n".join(ocr_text_parts)
        logger.info(f"Total OCR text extracted: {len(full_ocr_text)} characters")
        return full_ocr_text

    def learn_from_example(
        self,
        processor_id: str,
        input_file_bytes: bytes,
        desired_output: str,
        filename: str = None
    ) -> dict:
        """
        Learn transformation from an example file (any supported type).

        Supports: PDF, Word (.docx, .doc), Excel (.xlsx, .xls), CSV, TXT, Images (.png, .jpg, .jpeg)

        Stores input images (for structure) + text (for accuracy) + desired output.

        Generic approach - works for ANY document type.

        Args:
            processor_id: Unique ID for this processor
            input_file_bytes: Example file bytes
            desired_output: Desired formatted output
            filename: Optional filename for type detection

        Returns:
            Dict with success status and stats
        """
        logger.info(f"Learning transformation for processor '{processor_id}'")

        # Extract content from file (auto-detects type)
        input_images, text_content, file_type = self.extract_content_from_file(
            input_file_bytes,
            filename
        )

        logger.info(f"Detected file type: {file_type}, {len(input_images)} images, {len(text_content)} chars text")

        # Store the example
        self.examples[processor_id] = {
            'input_images': input_images,
            'ocr_text': text_content,
            'output_text': desired_output,
            'source_type': file_type  # Track source type (pdf, docx, csv, etc.)
        }

        logger.info(f"Learned example: {len(input_images)} images, {len(text_content)} chars text, {len(desired_output)} chars output")

        return {
            'success': True,
            'processor_id': processor_id,
            'input_length': len(input_images) if input_images else len(text_content),
            'file_type': file_type,
            'output_length': len(desired_output)
        }

    def learn_from_text(
        self,
        processor_id: str,
        input_text: str,
        desired_output: str
    ) -> dict:
        """
        Learn transformation from pasted text (no PDF).

        Stores input text + desired output (no images or OCR).
        For text-based sources like hockey stats, website copy/paste, etc.

        Generic approach - works for ANY document type.

        Args:
            processor_id: Unique ID for this processor
            input_text: Example input text (pasted)
            desired_output: Desired formatted output

        Returns:
            Dict with success status and stats
        """
        logger.info(f"Learning transformation for processor '{processor_id}' (text input)")

        # Store the example (no images - indicates text source)
        self.examples[processor_id] = {
            'input_images': [],  # Empty - no PDF
            'ocr_text': input_text,  # Use input text directly
            'output_text': desired_output,
            'source_type': 'text'  # Track source type
        }

        logger.info(f"Learned text example: {len(input_text)} chars input, {len(desired_output)} chars output")

        return {
            'success': True,
            'processor_id': processor_id,
            'input_length': len(input_text),
            'output_length': len(desired_output)
        }

    def transform(
        self,
        processor_id: str,
        new_file_bytes: bytes,
        filename: str = None
    ) -> dict:
        """
        Transform a new file using learned example.

        Supports: PDF, Word (.docx, .doc), Excel (.xlsx, .xls), CSV, TXT, Images (.png, .jpg, .jpeg)

        Generic approach - works for ANY document type.

        Args:
            processor_id: ID of learned processor
            new_file_bytes: New file to transform
            filename: Optional filename for type detection

        Returns:
            Dict with transformed output
        """
        if processor_id not in self.examples:
            raise ValueError(f"Processor '{processor_id}' not found. Learn from example first.")

        logger.info(f"Transforming document with processor '{processor_id}'")

        # Get the example
        example = self.examples[processor_id]
        example_images = example['input_images']
        example_text = example['ocr_text']
        example_output = example['output_text']

        # Extract content from new file (auto-detects type)
        new_images, new_text, file_type = self.extract_content_from_file(
            new_file_bytes,
            filename
        )

        logger.info(f"Detected file type: {file_type}, {len(new_images)} images, {len(new_text)} chars text")

        # Build content for Claude based on whether we have images or text-only
        if new_images and example_images:
            # Vision + OCR mode (for PDFs and images)
            content = self._build_vision_ocr_content(
                example_images=example_images,
                example_ocr_text=example_text,
                example_output=example_output,
                new_images=new_images,
                new_ocr_text=new_text
            )
        else:
            # Text-only mode (for Word, Excel, CSV, TXT)
            content = self._build_text_only_prompt(
                example_text=example_text,
                example_output=example_output,
                new_text=new_text
            )

        # Call Claude API
        logger.info(f"Calling Claude for transformation ({file_type})...")
        response = self.client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            messages=[{
                "role": "user",
                "content": content
            }]
        )

        output_text = response.content[0].text

        logger.info(f"Transformation complete: {len(output_text)} chars output")

        return {
            'success': True,
            'output': output_text,
            'input_tokens': response.usage.input_tokens,
            'output_tokens': response.usage.output_tokens,
            'tokens_used': response.usage.input_tokens + response.usage.output_tokens
        }

    def transform_text_only(
        self,
        processor_id: str,
        new_text: str
    ) -> dict:
        """
        Transform pasted text using learned example (text-to-text, no vision).

        For pasted text input (e.g., hockey stats), use only text transformation.
        Skips OCR and vision processing.

        Generic approach - works for ANY document type.

        Args:
            processor_id: ID of learned processor
            new_text: Pasted text to transform

        Returns:
            Dict with transformed output
        """
        if processor_id not in self.examples:
            raise ValueError(f"Processor '{processor_id}' not found. Learn from example first.")

        logger.info(f"Transforming text with processor '{processor_id}'")

        # Get the example
        example = self.examples[processor_id]
        example_ocr_text = example['ocr_text']
        example_output = example['output_text']

        # Build text-only prompt (no images)
        prompt = self._build_text_only_prompt(
            example_text=example_ocr_text,
            example_output=example_output,
            new_text=new_text
        )

        # Call Claude API (text only, no vision)
        logger.info("Calling Claude for text transformation...")
        response = self.client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            messages=[{
                "role": "user",
                "content": prompt
            }]
        )

        output_text = response.content[0].text

        logger.info(f"Text transformation complete: {len(output_text)} chars output")

        return {
            'success': True,
            'output': output_text,
            'input_tokens': response.usage.input_tokens,
            'output_tokens': response.usage.output_tokens,
            'tokens_used': response.usage.input_tokens + response.usage.output_tokens
        }

    def _build_text_only_prompt(
        self,
        example_text: str,
        example_output: str,
        new_text: str
    ) -> str:
        """
        Build text-only prompt for Claude API (no vision).

        For pasted text input - simpler and faster than vision.
        Generic approach - works for ANY document type.

        Args:
            example_text: OCR text from example
            example_output: Desired output format
            new_text: New text to transform

        Returns:
            Prompt string
        """
        return f"""You are a document transformer. Your job is to transform documents in a consistent format.

I will show you an EXAMPLE transformation, then give you NEW text to transform the same way.

Generic approach - works for honor rolls, legal notices, sports stats, or ANY document type.

# EXAMPLE TRANSFORMATION

## Example Input Text:
```
{example_text}
```

## Example Output (desired format):
```
{example_output}
```

# YOUR TASK

Transform this NEW text in the SAME WAY as the example above.

## New Input Text:
```
{new_text}
```

# INSTRUCTIONS

IMPORTANT RULES (generic - apply to ANY document type):

1. **Include ALL items, even zeros** - If an item has all zero values (0-0-0), still include it in the output. Do not drop items just because they have zero values.

2. **Detect threshold values** - Look for values that meet specific thresholds (e.g., if the example shows "Fouled out" for certain values, detect when new items meet that same threshold). Check the example to see how threshold detection is demonstrated.

3. **Complete extraction** - Extract every item from the document, not just the ones with non-zero values. Completeness is more important than brevity.

Provide ONLY the transformed output in the same format as the example.
Extract ALL details accurately - names, numbers, statistics, etc.
Do not include any explanation or commentary."""

    def _build_vision_ocr_content(
        self,
        example_images: List[str],
        example_ocr_text: str,
        example_output: str,
        new_images: List[str],
        new_ocr_text: str
    ) -> list:
        """
        Build vision + OCR content for Claude API.

        Sends Claude BOTH:
        - Images (for understanding structure/layout)
        - OCR text (for accurate names/values)

        Generic approach - works for ANY document type.

        Args:
            example_images: Base64-encoded images from example PDF
            example_ocr_text: Accurate OCR text from example PDF
            example_output: Desired output format
            new_images: Base64-encoded images from new PDF
            new_ocr_text: Accurate OCR text from new PDF

        Returns:
            List of content blocks for Claude API
        """
        content = []

        # Add instruction text
        content.append({
            "type": "text",
            "text": """You are a document transformer. Your job is to transform documents in a consistent format.

I will show you an EXAMPLE transformation, then give you a NEW document to transform the same way.

For each document, you receive:
1. IMAGES - to understand the structure and layout
2. OCR TEXT - for accurate names, numbers, and values (use this for precise extraction)

Generic approach - works for honor rolls, legal notices, sports stats, or ANY document type.

# EXAMPLE TRANSFORMATION

## Example Images (for structure/layout):"""
        })

        # Add example images
        for i, img_b64 in enumerate(example_images):
            content.append({
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": "image/png",
                    "data": img_b64
                }
            })

        # Add example OCR text
        content.append({
            "type": "text",
            "text": f"""
## Example OCR Text (accurate text extraction):
```
{example_ocr_text}
```

## Example Output (desired format):
```
{example_output}
```

# YOUR TASK

Transform this NEW document in the SAME WAY as the example above.

## New Document Images (for structure/layout):"""
        })

        # Add new images to transform
        for i, img_b64 in enumerate(new_images):
            content.append({
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": "image/png",
                    "data": img_b64
                }
            })

        # Add new OCR text
        content.append({
            "type": "text",
            "text": f"""
## New Document OCR Text (accurate text extraction):
```
{new_ocr_text}
```

# INSTRUCTIONS

Use the IMAGES to understand structure and layout.
Use the OCR TEXT for accurate names, numbers, and values.

**IMPORTANT:** You receive both images and OCR text. OCR text extraction can sometimes fail on complex PDFs - names may appear separated from their data, or rows may be out of order. Always cross-reference the OCR text with what you SEE in the images. If they conflict, trust the visual data from the images.

IMPORTANT RULES (generic - apply to ANY document type):

1. **Include ALL items, even zeros** - If an item has all zero values (0-0-0), still include it in the output. Do not drop items just because they have zero values.

2. **Detect threshold values** - Look for values that meet specific thresholds (e.g., if the example shows "Fouled out" for certain values, detect when new items meet that same threshold). Check the example to see how threshold detection is demonstrated.

3. **Complete extraction** - Extract every item from the document, not just the ones with non-zero values. Completeness is more important than brevity.

Provide ONLY the transformed output in the same format as the example.
Extract ALL details accurately - names, numbers, statistics, etc.
Do not include any explanation or commentary."""
        })

        return content


class SimpleTransformerDB:
    """
    Version of SimpleTransformer that persists examples to database.

    This allows examples to persist across sessions.
    """

    def __init__(self, db, api_key: Optional[str] = None):
        """
        Initialize with database connection.

        Args:
            db: Database instance
            api_key: Anthropic API key
        """
        self.db = db
        self.transformer = SimpleTransformer(api_key=api_key)

    async def learn_from_example(
        self,
        processor_id: str,
        name: str,
        input_file_bytes: bytes,
        desired_output: str,
        user_id: Optional[str] = None,
        filename: str = None
    ) -> dict:
        """Learn from any file type and save to database."""
        # Learn using simple transformer
        result = self.transformer.learn_from_example(
            processor_id=processor_id,
            input_file_bytes=input_file_bytes,
            desired_output=desired_output,
            filename=filename
        )

        # Save example to database
        example = self.transformer.examples[processor_id]

        # Store in database using existing schema
        # We'll store it as a special "simple" processor type
        from src.processors.models import Processor
        import json

        # Store images, OCR text, output, and source type in template field as JSON
        template_data = {
            'input_images': example['input_images'],
            'ocr_text': example['ocr_text'],
            'output_text': example['output_text'],
            'source_type': example.get('source_type', 'pdf')  # Default to PDF
        }

        processor = Processor(
            id=processor_id,
            name=name,
            document_type="simple_transform_vision_ocr",
            anchors=[],
            regions=[],
            extraction_ops=[],
            template_id="simple_vision_ocr",
            template=json.dumps(template_data)
        )

        processor_json = processor.to_json()

        await self.db.create_processor(
            processor_id=processor_id,
            name=name,
            document_type="simple_transform_vision_ocr",
            processor_json=processor_json,
            user_id=user_id
        )

        logger.info(f"Saved simple transformer processor '{name}' (PDF) to database")

        return result

    async def learn_from_text(
        self,
        processor_id: str,
        name: str,
        input_text: str,
        desired_output: str,
        user_id: Optional[str] = None
    ) -> dict:
        """Learn from text and save to database."""
        # Learn using simple transformer
        result = self.transformer.learn_from_text(
            processor_id=processor_id,
            input_text=input_text,
            desired_output=desired_output
        )

        # Save example to database
        example = self.transformer.examples[processor_id]

        # Store in database using existing schema
        from src.processors.models import Processor
        import json

        # Store text, output, and source type in template field as JSON
        template_data = {
            'input_images': [],  # Empty for text sources
            'ocr_text': example['ocr_text'],
            'output_text': example['output_text'],
            'source_type': example.get('source_type', 'text')  # Text source
        }

        processor = Processor(
            id=processor_id,
            name=name,
            document_type="simple_transform_text",  # Different type for text-based
            anchors=[],
            regions=[],
            extraction_ops=[],
            template_id="simple_text",
            template=json.dumps(template_data)
        )

        processor_json = processor.to_json()

        await self.db.create_processor(
            processor_id=processor_id,
            name=name,
            document_type="simple_transform_text",
            processor_json=processor_json,
            user_id=user_id
        )

        logger.info(f"Saved simple transformer processor '{name}' (text) to database")

        return result

    async def transform(
        self,
        processor_id: str,
        new_file_bytes: bytes,
        filename: str = None
    ) -> dict:
        """Transform using saved processor (supports all file types)."""
        # Load example from database if not in memory
        if processor_id not in self.transformer.examples:
            processor_data = await self.db.get_processor(processor_id)

            if not processor_data:
                raise ValueError(f"Processor '{processor_id}' not found")

            from src.processors.models import Processor
            import json

            processor = Processor.from_json(processor_data['processor_json'])

            # Parse example from template field
            template = processor.template

            # Try to parse as JSON (vision + OCR format)
            try:
                template_data = json.loads(template)

                # New format with OCR text and source type
                if 'input_images' in template_data and 'ocr_text' in template_data and 'output_text' in template_data:
                    self.transformer.examples[processor_id] = {
                        'input_images': template_data['input_images'],
                        'ocr_text': template_data['ocr_text'],
                        'output_text': template_data['output_text'],
                        'source_type': template_data.get('source_type', 'pdf')  # Default to PDF for old processors
                    }
                # Old vision-only format (backwards compatibility)
                elif 'input_images' in template_data and 'output_text' in template_data:
                    # No OCR text stored - will need to regenerate or error
                    raise ValueError(f"Processor '{processor_id}' uses old vision-only format. Please recreate with OCR support.")
            except json.JSONDecodeError:
                # Very old text-based format (backwards compatibility)
                if "EXAMPLE_INPUT:" in template and "EXAMPLE_OUTPUT:" in template:
                    raise ValueError(f"Processor '{processor_id}' uses old text format. Please recreate with vision + OCR support.")

        # Transform
        return self.transformer.transform(processor_id, new_file_bytes, filename)

    async def transform_text(
        self,
        processor_id: str,
        new_text: str
    ) -> dict:
        """
        Transform pasted text using saved processor.

        For text input (e.g., hockey stats), skip OCR and use text directly.
        """
        # Load example from database if not in memory
        if processor_id not in self.transformer.examples:
            processor_data = await self.db.get_processor(processor_id)

            if not processor_data:
                raise ValueError(f"Processor '{processor_id}' not found")

            from src.processors.models import Processor
            import json

            processor = Processor.from_json(processor_data['processor_json'])

            # Parse example from template field
            template = processor.template

            # Try to parse as JSON (vision + OCR format)
            try:
                template_data = json.loads(template)

                # New format with OCR text and source type
                if 'input_images' in template_data and 'ocr_text' in template_data and 'output_text' in template_data:
                    self.transformer.examples[processor_id] = {
                        'input_images': template_data['input_images'],
                        'ocr_text': template_data['ocr_text'],
                        'output_text': template_data['output_text'],
                        'source_type': template_data.get('source_type', 'text')  # Default to text for transform_text
                    }
                # Old vision-only format (backwards compatibility)
                elif 'input_images' in template_data and 'output_text' in template_data:
                    raise ValueError(f"Processor '{processor_id}' uses old vision-only format. Please recreate with OCR support.")
            except json.JSONDecodeError:
                # Very old text-based format (backwards compatibility)
                if "EXAMPLE_INPUT:" in template and "EXAMPLE_OUTPUT:" in template:
                    raise ValueError(f"Processor '{processor_id}' uses old text format. Please recreate with vision + OCR support.")

        # Transform using text input
        return self.transformer.transform_text_only(processor_id, new_text)
